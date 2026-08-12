import sys
import os
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'ORBIT-I'))
sys.path.insert(0, BASE_DIR)

import streamlit as st
import time
import pdfplumber
import docx

from classifier.domain_classifier import classify_resume
from classifier.preprocess import preprocess_text
from classifier.keyword_matcher import keyword_match
from classifier.confidence_score import calculate_confidence
from services.audit_logger import log_event


def extract_candidate_name(text):
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'objective',
                  'summary', 'profile', 'contact', 'email', 'phone',
                  'address', 'linkedin', 'github', 'dear', 'sir', 'madam']
    for line in lines[:10]:
        if any(char in line for char in ['@', 'http', 'www', '+92', '0300', '/', '📧', '📞', '📍']):
            continue
        if any(word in line.lower() for word in skip_words):
            continue
        if len(line) > 50:
            continue
        if any(char in line for char in ['|', '•', '·', '─', '=', ':', ',']):
            continue
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(word[0].isupper() for word in words if word.isalpha()):
                return line
    return "Candidate"


def get_position_title(domain):
    domain_lower = domain.lower()
    position_mapping = {
        "software engineering": "Software Engineer",
        "software development": "Software Developer",
        "web development": "Web Developer",
        "frontend development": "Frontend Developer",
        "backend development": "Backend Developer",
        "full stack development": "Full Stack Developer",
        "mobile development": "Mobile App Developer",
        "data science": "Data Scientist",
        "data analysis": "Data Analyst",
        "data engineering": "Data Engineer",
        "machine learning": "Machine Learning Engineer",
        "artificial intelligence": "AI Engineer",
        "cybersecurity": "Cybersecurity Analyst",
        "cyber security": "Cybersecurity Analyst",
        "network security": "Network Security Engineer",
        "ui/ux design": "UI/UX Designer",
        "ui ux design": "UI/UX Designer",
        "graphic design": "Graphic Designer",
        "product design": "Product Designer",
        "cloud computing": "Cloud Engineer",
        "devops": "DevOps Engineer",
        "business analysis": "Business Analyst",
        "project management": "Project Manager",
        "quality assurance": "QA Engineer",
        "database administration": "Database Administrator",
        "public health": "Public Health Officer",
        "finance": "Financial Analyst",
        "marketing": "Marketing Specialist",
        "human resources": "HR Executive",
        "sales": "Sales Executive",
    }
    for key, value in position_mapping.items():
        if key in domain_lower:
            return value
    return f"{domain} Professional"


def get_all_domain_scores(resume_text):
    """Return confidence scores for every domain so the user sees the full picture."""
    preprocessed = preprocess_text(resume_text)
    match_results = keyword_match(preprocessed)
    confidence_results = calculate_confidence(match_results)
    scores = {
        domain: info["confidence"]
        for domain, info in confidence_results.items()
    }
    return dict(sorted(scores.items(), key=lambda x: x[1], reverse=True))


# ─────────────────────────────────────────────
st.set_page_config(page_title="ORBIT-I | Upload", page_icon="📂", layout="wide")

with open("assets/style.css", encoding="utf-8") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

st.markdown("""
    <style>
        [data-testid="stSidebarNav"] ul li:first-child {display: none;}
        [data-testid="stSidebarNav"]::before {
            content: "ORBIT-I";
            display: block;
            font-size: 20px;
            font-weight: 700;
            color: var(--lb-primary, #1a73e8);
            padding: 24px 16px 16px 16px;
            letter-spacing: 1px;
        }
        [data-testid="stMetricValue"] { font-size: 16px !important; }
        [data-testid="stMetricLabel"] { font-size: 12px !important; }
    </style>
""", unsafe_allow_html=True)

for key, default in [
    ("total_uploaded", 0), ("pending", 0),
    ("processed", 0), ("last_uploaded", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default

col_title, col_top = st.columns([3, 2])
with col_title:
    st.markdown("""
        <div style="display:flex; align-items:center; gap:12px;">
            <div class="stat-icon-circle">📂</div>
            <div>
                <h1 style="margin:0; padding:0; line-height:1.1;">Upload Resume</h1>
                <p style="margin:0; color: var(--lb-text-muted); font-size:14px;">Upload candidate CVs in PDF or DOCX format for AI processing.</p>
            </div>
        </div>
    """, unsafe_allow_html=True)
with col_top:
    st.markdown("""
        <div class="orbit-topbar">
            <div class="orbit-search">🔍 <span>Search...</span></div>
            <div class="orbit-icon-btn">🔔</div>
            <div class="orbit-avatar">A</div>
        </div>
    """, unsafe_allow_html=True)

st.divider()

uploaded_file = st.file_uploader(
    "Drag & Drop your CV here — OR click Browse files",
    type=["pdf", "docx"],
)
st.caption("Supported formats: PDF, DOCX  •  Max file size: 20 MB")

if uploaded_file is not None:

    is_new_file = st.session_state.last_uploaded != uploaded_file.name
    if is_new_file:
        st.session_state.last_uploaded = uploaded_file.name
        st.session_state.total_uploaded += 1
        st.session_state.pending += 1

    file_size_kb = len(uploaded_file.getvalue()) / 1024
    file_size_label = f"{file_size_kb/1024:.1f} MB" if file_size_kb > 1024 else f"{file_size_kb:.0f} KB"
    file_icon = "📕" if uploaded_file.name.lower().endswith(".pdf") else "📘"

    st.markdown(f"""
        <div class="file-row">
            <span class="file-icon">{file_icon}</span>
            <span class="file-name">{uploaded_file.name}</span>
            <span class="file-size">{file_size_label}</span>
            <span class="status-pill uploaded">Uploaded</span>
        </div>
    """, unsafe_allow_html=True)

    with st.spinner("Processing your CV, please wait..."):

        extracted_text = ""

        if uploaded_file.name.endswith(".pdf"):
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
                    for table in page.extract_tables():
                        for row in table:
                            for cell in row:
                                if cell:
                                    extracted_text += cell + " "

        elif uploaded_file.name.endswith(".docx"):
            document = docx.Document(uploaded_file)
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + "\n"

        candidate_name = extract_candidate_name(extracted_text)

        result = None
        offer_result = None
        domain = "Unknown"
        confidence = 0
        status = "Manual Review"
        all_scores = {}

        if extracted_text.strip():
            result = classify_resume(extracted_text)
            all_scores = get_all_domain_scores(extracted_text)

            if st.session_state.get("last_processed") != uploaded_file.name:
                st.session_state.last_processed = uploaded_file.name
                st.session_state.processed += 1
                if st.session_state.pending > 0:
                    st.session_state.pending -= 1

            domain     = result.get("predicted_domain", "Unknown")
            confidence = result.get("confidence", 0)
            status     = result.get("status", "Manual Review")

            # Always store full classification detail in session so
            # Manual Override can access per-domain scores
            st.session_state["all_domain_scores"] = all_scores
            st.session_state["extracted_text"]    = extracted_text

            if confidence >= 75:
                position_title = get_position_title(domain)
                st.session_state["candidate_data"] = {
                    "name":         candidate_name,
                    "email":        "",
                    "phone":        "",
                    "position":     position_title,
                    "salary":       "100000",
                    "joining_date": "",
                    "domain":       domain,
                    "confidence":   confidence,
                    "remarks":      "",
                }
                offer_result = {"success": True}
            else:
                # Still store data so Manual Override can pre-populate
                st.session_state["candidate_data"] = {
                    "name":         candidate_name,
                    "email":        "",
                    "phone":        "",
                    "position":     get_position_title(domain),
                    "salary":       "100000",
                    "joining_date": "",
                    "domain":       domain,
                    "confidence":   confidence,
                    "remarks":      "",
                }

            log_event(
                cv_filename=uploaded_file.name,
                domain_assigned=domain,
                confidence_score=confidence,
                offer_status="Generated" if offer_result and offer_result.get("success") else "Flagged for Review",
                edited_by="System",
                notes=f"Confidence: {confidence}% | Candidate: {candidate_name}",
            )

        time.sleep(1)

    st.success(f"✅ File uploaded successfully: {uploaded_file.name}")

    # ── AI Validation checklist ──
    text_len = len(extracted_text.strip())
    file_ext_valid = uploaded_file.name.lower().endswith((".pdf", ".docx"))
    text_extracted_ok = text_len > 200
    low_quality = 0 < text_len <= 200

    with st.container(border=True):
        st.markdown("##### AI Validation")

        v_icon = "✅" if file_ext_valid else "❌"
        st.markdown(f"""
            <div class="validation-item">
                <span class="v-icon">{v_icon}</span>
                <div>
                    <div class="v-title">File Format Valid</div>
                    <div class="v-desc">{uploaded_file.name.split('.')[-1].upper()} format is supported</div>
                </div>
            </div>
        """, unsafe_allow_html=True)

        v_icon = "✅" if text_extracted_ok else "⚠️"
        st.markdown(f"""
            <div class="validation-item">
                <span class="v-icon">{v_icon}</span>
                <div>
                    <div class="v-title">Text Extraction</div>
                    <div class="v-desc">{"Content extracted successfully" if text_extracted_ok else "Little to no readable text found"}</div>
                </div>
            </div>
        """, unsafe_allow_html=True)

        if low_quality or not text_len:
            st.markdown("""
                <div class="validation-item">
                    <span class="v-icon">⚠️</span>
                    <div>
                        <div class="v-title">Low Scan Quality</div>
                        <div class="v-desc">Some sections may have low clarity — consider a text-based PDF or DOCX</div>
                    </div>
                </div>
            """, unsafe_allow_html=True)

    st.divider()

    if result:
        st.subheader("🎯 Classification Result")

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Candidate", candidate_name)
        with col2:
            st.metric("Auto-Detected Domain", domain)
        with col3:
            st.metric("Confidence Score", f"{confidence}%")
        with col4:
            if status == "Manual Review":
                st.metric("Status", "⚠️ Manual Review")
            else:
                st.metric("Status", "✅ Auto Classified")

        # ── All Domain Scores breakdown ──────────────────────────────
        if all_scores:
            st.divider()
            st.markdown("##### 📊 All Domain Scores")
            st.caption("Shows how strongly your CV matches every domain in the system.")

            score_cols = st.columns(min(len(all_scores), 4))
            for i, (d, s) in enumerate(all_scores.items()):
                col = score_cols[i % len(score_cols)]
                icon = "🥇" if d == domain else ""
                col.metric(f"{icon} {d}", f"{s}%")

        st.divider()

        if status == "Manual Review":
            st.warning(
                f"⚠️ Auto-classification confidence ({confidence}%) is below 75%. "
                "Please go to Manual Override to select the correct domain."
            )
            if st.button("✏️ Go to Manual Override"):
                st.switch_page("pages/manual_override.py")
        else:
            if offer_result and offer_result.get("success"):
                st.success("✅ Candidate information extracted successfully.")
                if st.button("➡ Continue to Manual Override", use_container_width=True):
                    st.switch_page("pages/manual_override.py")
            elif offer_result and offer_result.get("error"):
                st.error(f"❌ Offer generation failed: {offer_result.get('error')}")

    st.divider()
    if st.button("🏠 Back to Home"):
        st.switch_page("pages/0_home.py")
