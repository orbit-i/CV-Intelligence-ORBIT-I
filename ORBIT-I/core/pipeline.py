import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from classifier.domain_classifier import classify_resume
from core.offer_generator import generate_offer
from services.audit_logger import log_event


def run_pipeline(uploaded_file, candidate_name="", salary="PKR 100,000 / month"):
    """
    Full pipeline:
    1. Text extract
    2. Classify domain
    3. Generate offer letter (if score >= 75)
    4. Log to audit
    """

    result = {
        "candidate_name": candidate_name,
        "domain": None,
        "confidence": 0,
        "status": None,
        "offer_letter": None,
        "error": None
    }

    try:
        # Step 1: Text extract
        extracted_text = ""

        if hasattr(uploaded_file, 'name'):
            filename = uploaded_file.name
        else:
            filename = str(uploaded_file)

        if filename.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            for cell in row:
                                if cell:
                                    extracted_text += cell + " "

        elif filename.endswith(".docx"):
            import docx
            document = docx.Document(uploaded_file)
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + "\n"

        if not extracted_text.strip():
            result["error"] = "Could not extract text from CV"
            return result

        # Step 2: Classify domain
        classification = classify_resume(extracted_text)
        domain = classification.get("predicted_domain", "Unknown")
        confidence = classification.get("confidence", 0)
        status = classification.get("status", "Manual Review")

        result["domain"] = domain
        result["confidence"] = confidence
        result["status"] = status

        # Step 3: Generate offer letter if score >= 75
        if confidence >= 75:
            candidate_profile = {
                "candidate_name": candidate_name or "Candidate",
                "domain": domain,
                "position_title": domain,
                "salary": salary,
                "company_name": "ORBIT-I",
                "hr_signatory": "HR Department",
                "probation_period": "3 months",
                "location": "Hybrid - Karachi, Pakistan",
            }

            offer_result = generate_offer(candidate_profile)

            if "error" in offer_result:
                result["error"] = offer_result["error"]
            else:
                result["offer_letter"] = offer_result["offer_letter"]

        # Step 4: Log to audit
        log_event(
            cv_filename=filename,
            domain_assigned=domain,
            confidence_score=confidence,
            offer_status="Generated" if result["offer_letter"] else "Flagged for Review",
            edited_by="System",
            notes=f"Confidence: {confidence}%"
        )

    except Exception as e:
        result["error"] = str(e)

    return result