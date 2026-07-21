from docx import Document
from pathlib import Path
from datetime import datetime, timedelta


def replace_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)


def generate_offer(candidate_profile):
    base_dir = Path(__file__).resolve().parent.parent

    template_path = base_dir / "templates" / "offer_template.docx"

    if not template_path.exists():
        return {"error": f"Template not found at: {template_path}"}

    doc = Document(template_path)

    offer_date = datetime.now().strftime("%B %d, %Y")
    start_date = candidate_profile.get(
        "start_date",
        (datetime.now() + timedelta(weeks=2)).strftime("%B %d, %Y")
    )

    replacements = {
        "{{candidate_name}}": candidate_profile.get("candidate_name", ""),
        "{{domain}}":         candidate_profile.get("domain", ""),
        "{{position_title}}": candidate_profile.get("position_title", candidate_profile.get("domain", "")),
        "{{salary}}":         candidate_profile.get("salary", ""),
        "{{start_date}}":     start_date,
        "{{company_name}}":   candidate_profile.get("company_name", "iCompany Pakistan"),
        "{{hr_signatory}}":   candidate_profile.get("hr_signatory", "HR Department"),
        "{{offer_date}}":     offer_date,
        "{{probation_period}}": candidate_profile.get("probation_period", "3 months"),
        "{{location}}":       candidate_profile.get("location", "Hybrid - Karachi, Pakistan"),
    }

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    output_dir = base_dir / "data" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    candidate_name = candidate_profile.get("candidate_name", "candidate").replace(" ", "_")
    output_file = output_dir / f"{candidate_name}_offer.docx"

    if output_file.exists():
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"{candidate_name}_offer_{timestamp}.docx"

    doc.save(output_file)

    return {
        "success": True,
        "offer_letter": str(output_file),
        "candidate_name": candidate_profile.get("candidate_name", ""),
        "domain": candidate_profile.get("domain", ""),
    }
