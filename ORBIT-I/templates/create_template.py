from docx import Document
import os

os.makedirs('templates', exist_ok=True)

doc = Document()
doc.add_heading('Offer Letter', 0)

doc.add_paragraph('Date: {{today_date}}')
doc.add_paragraph('')
doc.add_paragraph('Dear {{candidate_name}},')
doc.add_paragraph('')
doc.add_paragraph('We are pleased to offer you the position of {{domain}} at our company.')
doc.add_paragraph('Your monthly compensation will be PKR {{salary}}.')
doc.add_paragraph('Your joining date is {{start_date}}. You will be based in {{location}}.')
doc.add_paragraph('')
doc.add_paragraph('Sincerely,')
doc.add_paragraph('HR Department')

doc.save('templates/offer_letter_template.docx')
print("Template ban gaya: templates/offer_letter_template.docx")
Where to paste