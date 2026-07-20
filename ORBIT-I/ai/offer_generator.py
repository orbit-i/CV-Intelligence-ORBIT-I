from docx import Document
from pathlib import Path
from datetime import datetime, timedelta


def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in paragraph runs to preserve formatting."""
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)


def generate_offer(candidate_profile):
    """
    Generate an offer letter DOCX from template.

    candidate_profile example:
    {
        "candidate_name": "Ali Khan",
        "domain": "Software Engineering",
        "position_title": "Software Engineer",
        "salary": "PKR 100,000 / month",
        "start_date": "2026-08-01",
        "company_name": "iCompany Pakistan",
        "hr_signatory": "HR Department",
        "probation_period": "3 months",
        "location": "Hybrid - Karachi, Pakistan",
    }
    """

    base_dir = Path(__file__).resolve().parent.parent

    template_path = base_dir / "templates" / "offer_template.docx"

    if not template_path.exists():
        return {"error": f"Template not found at: {template_path}"}

    doc = Document(template_path)

    # Auto generate offer date and start date if not provided
    offer_date = datetime.now().strftime("%B %d, %Y")
    start_date = candidate_profile.get(
        "start_date",
        (datetime.now() + timedelta(weeks=2)).strftime("%B %d, %Y")
    )

    replacements = {
        "{{candidate_name}}": candidate_profile.get("candidate_name", ""),
        "{{domain}}":         candidate_profile.get("domain", ""),
        "{{position_title}}": candidate_profile.get("position_title", candidate_profile.get("domain", "")),
        "{{salary}}":         candidate_profile.get("salary", ""),
        "{{start_date}}":     start_date,
        "{{company_name}}":   candidate_profile.get("company_name", "iCompany Pakistan"),
        "{{hr_signatory}}":   candidate_profile.get("hr_signatory", "HR Department"),
        "{{offer_date}}":     offer_date,
        "{{probation_period}}": candidate_profile.get("probation_period", "3 months"),
        "{{location}}":       candidate_profile.get("location", "Hybrid - Karachi, Pakistan"),
    }

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # Replace inside tables (some templates use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    # Save output to data/output folder
    output_dir = base_dir / "data" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    candidate_name = candidate_profile.get("candidate_name", "candidate").replace(" ", "_")
    output_file = output_dir / f"{candidate_name}_offer.docx"

    # Avoid overwriting existing file
    if output_file.exists():
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"{candidate_name}_offer_{timestamp}.docx"

    doc.save(output_file)

    return {
        "success": True,
        "offer_letter": str(output_file),
        "candidate_name": candidate_profile.get("candidate_name", ""),
        "domain": candidate_profile.get("domain", ""),
    }


if __name__ == "__main__":
    test_data = {
        "candidate_name": "Mahnoor Gill",
        "domain": "Software Engineering",
        "position_title": "Software Engineer",
        "salary": "PKR 100,000 / month",
        "start_date": "August 01, 2026",
        "company_name": "iCompany Pakistan",
        "hr_signatory": "HR Department",
        "probation_period": "3 months",
        "location": "Hybrid - Karachi, Pakistan",
    }

    result = generate_offer(test_data)

    if "error" in result:
        print("Error:", result["error"])
    else:
        print("Success! Offer saved to:", result["offer_letter"])
